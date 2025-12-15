#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
一次性工具：把“编号清晰”的访谈提纲 docx 转成结构化 questions.json + questions.txt

用法：
  python tools/convert_docx_questions.py "survey/xxx.docx"
  python tools/convert_docx_questions.py "survey/xxx.docx" --out-dir survey

输出：
  <out-dir>/questions.json
  <out-dir>/questions.txt
"""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document


HEADING_RE = re.compile(r"^\s*[一二三四五六七八九十]+、\s*(.+?)（\s*\d+\s*题\s*）\s*$")


def _cell_text(cell) -> str:
    return "\n".join((p.text or "").strip() for p in cell.paragraphs if (p.text or "").strip()).strip()


def extract_categories(docx_path: Path) -> list[dict]:
    """
    按 docx 中的“四类标题 + 四个表格”输出分类问题：
    [
      { "title": "三、学龄前康复阶段（17 题）", "questions": [ {id,text}, ... ] },
      ...
    ]
    """
    doc = Document(str(docx_path))

    headings: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        if HEADING_RE.match(t):
            headings.append(t)

    # 模板：标题数与表格数一致时，按顺序对应；否则兜底按表格顺序生成
    titles = headings if headings and len(headings) == len(doc.tables) else [f"问题类别 {i+1}" for i in range(len(doc.tables))]

    categories: list[dict] = []
    for ti, tb in enumerate(doc.tables):
        title = titles[ti] if ti < len(titles) else f"问题类别 {ti+1}"
        qs: list[dict] = []

        for row in tb.rows:
            if len(row.cells) < 2:
                continue
            c0 = _cell_text(row.cells[0])
            c1 = _cell_text(row.cells[1])
            if not c0 or not c1:
                continue
            if c0 in ("题号", "编号", "序号"):
                continue
            if c1 in ("提纲原题",):
                continue
            if not re.fullmatch(r"\d{1,3}", c0):
                continue
            qs.append({"id": c0, "text": c1})

        categories.append({"title": title, "count": len(qs), "questions": qs})

    return categories


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", help="问题模板 docx 路径")
    ap.add_argument("--out-dir", default="survey", help="输出目录（默认 survey）")
    args = ap.parse_args()

    docx_path = Path(args.docx).expanduser().resolve()
    out_dir = Path(args.out_dir).expanduser().resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    categories = extract_categories(docx_path)
    total = sum(c.get("count", 0) for c in categories)
    payload = {
        "source_docx": str(docx_path),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "total_count": total,
        "categories": categories,
    }

    json_path = out_dir / "questions.json"
    txt_path = out_dir / "questions.txt"

    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    txt_lines: list[str] = []
    for cat in categories:
        txt_lines.append(f"### {cat['title']}")
        for q in cat.get("questions", []):
            txt_lines.append(f"{q['id']}. {q['text']}")
        txt_lines.append("")
    txt_path.write_text("\n".join(txt_lines).rstrip() + "\n", encoding="utf-8")

    print(f"OK: {total} questions")
    print(f"- {json_path}")
    print(f"- {txt_path}")


if __name__ == "__main__":
    main()

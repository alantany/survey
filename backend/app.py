import os
import uuid
import json
import time
import threading
import subprocess
import shutil
import re
import urllib.request
import urllib.error
import urllib.parse
import base64
import hashlib
import hmac
from pathlib import Path
from typing import Optional

from flask import Flask, jsonify, request, send_from_directory, send_file
from docx import Document

try:
    import requests
except ImportError:
    requests = None

try:
    import websocket
except ImportError:
    websocket = None


ROOT_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
WORK_DIR = DATA_DIR / "work"
SURVEY_DIR = ROOT_DIR / "survey"
CONFIG_PATH = ROOT_DIR / "config.json"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
WORK_DIR.mkdir(parents=True, exist_ok=True)
SURVEY_DIR.mkdir(parents=True, exist_ok=True)


def _env(name: str, default: str) -> str:
    v = os.environ.get(name)
    return v if v else default


def _first_existing_cmd(candidates: list[str], fallback: str) -> str:
    for c in candidates:
        if shutil.which(c):
            return c
    return fallback


# Homebrew 的 whisper-cpp 包通常提供 whisper-cli/whisper-server 等命令
WHISPER_BIN = os.environ.get("WHISPER_BIN") or _first_existing_cmd(
    ["whisper-cli", "whisper-cpp", "main"], "whisper-cli"
)
FFMPEG_BIN = os.environ.get("FFMPEG_BIN") or _first_existing_cmd(["ffmpeg"], "ffmpeg")
WHISPER_MODEL = _env("WHISPER_MODEL", str(ROOT_DIR / "models" / "ggml-small.bin"))
WHISPER_LANGUAGE = _env("WHISPER_LANGUAGE", "zh")
WHISPER_THREADS = int(_env("WHISPER_THREADS", str(min(8, os.cpu_count() or 4))))

# 允许更大文件（1小时音频可能很大）
MAX_CONTENT_LENGTH_MB = int(_env("MAX_CONTENT_LENGTH_MB", "1024"))  # 1GB


app = Flask(__name__, static_folder=None)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH_MB * 1024 * 1024


_jobs_lock = threading.Lock()
_jobs: dict[str, dict] = {}


def _load_local_config() -> dict:
    """
    读取本地 config.json（被 .gitignore 忽略）或环境变量。
    - OPENROUTER_API_KEY, OPENROUTER_MODEL, OPENROUTER_MAX_TOKENS
    - STT_API_URL, STT_API_KEY, STT_API_METHOD (stt_api_method: "POST"/"GET")
    """
    cfg: dict = {}
    try:
        if CONFIG_PATH.exists():
            cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        cfg = {}

    api_key = (cfg.get("openrouter_api_key") if isinstance(cfg, dict) else None) or os.environ.get("OPENROUTER_API_KEY")
    model = (cfg.get("openrouter_model") if isinstance(cfg, dict) else None) or os.environ.get("OPENROUTER_MODEL") or "tngtech/deepseek-r1t2-chimera:free"
    max_tokens = (cfg.get("openrouter_max_tokens") if isinstance(cfg, dict) else None) or os.environ.get("OPENROUTER_MAX_TOKENS") or 8192
    try:
        max_tokens = int(max_tokens)
    except Exception:
        max_tokens = 8192

    stt_api_url = (cfg.get("stt_api_url") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_URL") or ""
    stt_api_key = (cfg.get("stt_api_key") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_KEY") or ""
    stt_api_method = (cfg.get("stt_api_method") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_METHOD") or "POST"
    stt_api_type = (cfg.get("stt_api_type") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_TYPE") or ""
    stt_api_appid = (cfg.get("stt_api_appid") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_APPID") or ""
    stt_api_key = (cfg.get("stt_api_key") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_KEY") or ""
    stt_api_secret_key = (cfg.get("stt_api_secret_key") if isinstance(cfg, dict) else None) or os.environ.get("STT_API_SECRET_KEY") or ""

    return {
        "openrouter_api_key": (api_key or "").strip(),
        "openrouter_model": (model or "").strip(),
        "openrouter_max_tokens": max_tokens,
        "stt_api_url": (stt_api_url or "").strip(),
        "stt_api_key": (stt_api_key or "").strip(),
        "stt_api_method": (stt_api_method or "POST").strip().upper(),
        "stt_api_type": (stt_api_type or "").strip().lower(),
        "stt_api_appid": (stt_api_appid or "").strip(),
        "stt_api_secret_key": (stt_api_secret_key or "").strip(),
    }


def _strip_code_fence(s: str) -> str:
    """
    去掉常见的 ```json ... ``` 包裹，便于解析 JSON。
    """
    t = (s or "").strip()
    if t.startswith("```"):
        # ```json\n...\n``` 或 ```\n...\n```
        t = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", t)
        t = re.sub(r"\s*```$", "", t)
        t = t.strip()
    return t

def _safe_basename(name: str) -> str:
    """
    生成一个适合做文件名的 basename（尽量保留中文/英文/数字/_/-，其他替换为 _）
    """
    base = Path(name).stem.strip() or "audio"
    base = re.sub(r"[^\w\u4e00-\u9fff\-]+", "_", base)
    return base[:80] or "audio"


def _set_job(job_id: str, **kwargs):
    with _jobs_lock:
        job = _jobs.get(job_id, {})
        job.update(kwargs)
        _jobs[job_id] = job


def _get_job(job_id: str) -> Optional[dict]:
    with _jobs_lock:
        return _jobs.get(job_id)


def _run(cmd: list[str], cwd: Optional[Path] = None) -> tuple[int, str]:
    try:
        p = subprocess.Popen(
            cmd,
            cwd=str(cwd) if cwd else None,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            universal_newlines=True,
        )
        out_lines: list[str] = []
        assert p.stdout is not None
        for line in p.stdout:
            out_lines.append(line)
        rc = p.wait()
        return rc, "".join(out_lines)
    except FileNotFoundError as e:
        return 127, f"找不到命令：{e}\ncmd={cmd}\n"


def _run_stream(
    cmd: list[str],
    cwd: Optional[Path] = None,
    on_line=None,
    max_capture_lines: int = 5000,
) -> tuple[int, str]:
    """
    流式读取子进程输出，用于实时更新进度/日志。
    返回 (exit_code, captured_output)；captured_output 会被截断到 max_capture_lines 行。
    """
    try:
        p = subprocess.Popen(
            cmd,
            cwd=str(cwd) if cwd else None,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            universal_newlines=True,
        )
    except FileNotFoundError as e:
        return 127, f"找不到命令：{e}\ncmd={cmd}\n"

    out_lines: list[str] = []
    assert p.stdout is not None
    for line in p.stdout:
        if on_line:
            try:
                on_line(line)
            except Exception:
                pass
        out_lines.append(line)
        if len(out_lines) > max_capture_lines:
            out_lines = out_lines[-max_capture_lines:]
    rc = p.wait()
    return rc, "".join(out_lines)


def _get_audio_duration(audio_path: Path) -> int:
    """
    获取音频文件的时长（毫秒）
    使用 ffprobe 或 ffmpeg 获取时长
    """
    try:
        # 尝试使用 ffprobe
        ffprobe_bin = os.environ.get("FFPROBE_BIN") or _first_existing_cmd(["ffprobe"], "ffprobe")
        cmd = [
            ffprobe_bin,
            "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            str(audio_path)
        ]
        rc, out = _run(cmd)
        if rc == 0:
            duration_seconds = float(out.strip())
            return int(duration_seconds * 1000)  # 转换为毫秒
    except Exception:
        pass
    
    # 如果 ffprobe 失败，尝试使用 ffmpeg
    try:
        cmd = [
            FFMPEG_BIN,
            "-i", str(audio_path),
            "-f", "null",
            "-"
        ]
        rc, out = _run(cmd)
        # ffmpeg 输出中包含 Duration: HH:MM:SS.mmm
        import re
        match = re.search(r'Duration: (\d+):(\d+):(\d+\.\d+)', out)
        if match:
            hours = int(match.group(1))
            minutes = int(match.group(2))
            seconds = float(match.group(3))
            duration_seconds = hours * 3600 + minutes * 60 + seconds
            return int(duration_seconds * 1000)  # 转换为毫秒
    except Exception:
        pass
    
    # 如果都失败，返回默认值（200秒 = 200000毫秒）
    return 200000


def _to_wav_16k_mono(src: Path, dst: Path) -> tuple[bool, str]:
    cmd = [
        FFMPEG_BIN,
        "-y",
        "-i",
        str(src),
        "-ar",
        "16000",
        "-ac",
        "1",
        "-c:a",
        "pcm_s16le",
        str(dst),
    ]
    rc, out = _run(cmd)
    return rc == 0, out


def _whisper_transcribe(wav_path: Path, out_prefix: Path) -> tuple[bool, str]:
    # whisper.cpp 常见参数：-m 模型 -f 输入 -l 语言 -otxt 输出文本 -of 输出前缀
    cmd = [
        WHISPER_BIN,
        "-t",
        str(WHISPER_THREADS),
        "-m",
        WHISPER_MODEL,
        "-l",
        WHISPER_LANGUAGE,
        "-f",
        str(wav_path),
        "-pp",
        "-otxt",
        "-of",
        str(out_prefix),
    ]
    # 运行中实时抓进度：尽量兼容不同输出格式
    progress_re = re.compile(r"progress\\s*=\\s*(\\d+)%", re.IGNORECASE)
    any_percent_re = re.compile(r"(\\d{1,3})%")
    log_tail: list[str] = []
    last_progress: Optional[int] = None

    def on_line(line: str):
        nonlocal last_progress, log_tail
        s = line.strip()
        if not s:
            return

        log_tail.append(s)
        if len(log_tail) > 80:
            log_tail = log_tail[-80:]

        m = progress_re.search(s)
        if m:
            try:
                last_progress = int(m.group(1))
            except Exception:
                pass
        else:
            # 兜底：行里出现 xx% 就取一个
            m2 = any_percent_re.search(s)
            if m2:
                try:
                    v = int(m2.group(1))
                    if 0 <= v <= 100:
                        last_progress = v if last_progress is None else max(last_progress, v)
                except Exception:
                    pass

        # out_prefix.name == job_id（调用者传入 WORK_DIR/job_id）
        if last_progress is not None:
            _set_job(
                str(out_prefix.name),
                progress=last_progress,
                message=f"转写中… {last_progress}%",
                log_tail=log_tail,
            )
        else:
            _set_job(str(out_prefix.name), log_tail=log_tail)

    rc, out = _run_stream(cmd, cwd=ROOT_DIR, on_line=on_line)
    return rc == 0, out


def _xunfei_get_signa(appid: str, secret_key: str, ts: str) -> str:
    """
    科大讯飞旧版 API 签名生成（raasr.xfyun.cn）
    根据用户提供的 demo 代码实现
    """
    m2 = hashlib.md5()
    m2.update((appid + ts).encode('utf-8'))
    md5 = m2.hexdigest()
    md5 = bytes(md5, encoding='utf-8')
    signa = hmac.new(secret_key.encode('utf-8'), md5, hashlib.sha1).digest()
    signa = base64.b64encode(signa)
    return str(signa, 'utf-8')


def _xunfei_generate_signature(access_key_secret: str, params: dict) -> str:
    """
    科大讯飞官方文档签名生成（录音文件转写大模型接口）
    参考：https://www.xfyun.cn/doc/spark/asr_llm/Ifasr_llm.html
    
    1. 排除"signature"字段
    2. 按参数名自然排序
    3. URL编码每个key和value
    4. 拼接成 baseString: key1=value1&key2=value2
    5. HMAC-SHA1(baseString, access_key_secret)
    6. Base64编码
    """
    # 排除signature字段并排序
    sorted_params = {}
    for k, v in sorted(params.items()):
        if k != 'signature' and v is not None and str(v).strip():
            sorted_params[k] = str(v)
    
    # 构建baseString
    parts = []
    for k, v in sorted_params.items():
        # URL编码（标准URL编码，不保留特殊字符）
        encoded_key = urllib.parse.quote(k, safe='')
        encoded_value = urllib.parse.quote(v, safe='')
        parts.append(f"{encoded_key}={encoded_value}")
    
    base_string = '&'.join(parts)
    
    # HMAC-SHA1
    signature_bytes = hmac.new(
        access_key_secret.encode('utf-8'),
        base_string.encode('utf-8'),
        hashlib.sha1
    ).digest()
    
    # Base64编码
    signature = base64.b64encode(signature_bytes).decode('utf-8')
    return signature


def _xunfei_transcribe_new_api(audio_path: Path, job_id: str, appid: str, api_key: str, api_secret: str) -> tuple[bool, str]:
    """
    科大讯飞新版 HTTPS API 转写（office-api-ist-dx.iflyaisol.com）
    参考官方文档：https://www.xfyun.cn/doc/spark/asr_llm/Ifasr_llm.html
    
    返回 (成功, 文本或错误信息)
    """
    if requests is None:
        return False, "需要安装 requests 库：pip install requests"
    
    # HTTPS API 地址（根据官方文档）
    api_host = 'https://office-api-ist-dx.iflyaisol.com'
    api_upload = '/v2/upload'  # 正确的上传接口路径
    api_get_result = '/v2/getResult'  # 正确的结果查询接口路径
    
    try:
        from datetime import datetime
        import random
        import string
        
        # 根据官方文档，需要以下参数：
        # appId, accessKeyId, dateTime, signatureRandom, fileSize, fileName, language
        file_len = audio_path.stat().st_size
        file_name = audio_path.name
        
        # dateTime: 请求发起的本地时间，格式为 yyyy-MM-dd'T'HH:mm:ssZ
        # 根据错误信息，格式应该是 yyyy-MM-dd'T'HH:mm:ssZ，其中 Z 是时区偏移（如 +0800）
        now = datetime.now()
        timezone_offset = now.strftime('%z')
        if not timezone_offset:
            timezone_offset = '+0800'  # 默认东八区
        # 确保格式为 ±HHmm（例如 +0800）
        if len(timezone_offset) == 5:
            # 已经是 ±HHmm 格式
            pass
        elif len(timezone_offset) == 6:
            # 可能是 ±HH:MM 格式，需要去掉冒号
            timezone_offset = timezone_offset[:3] + timezone_offset[4:]
        date_time = now.strftime(f"%Y-%m-%dT%H:%M:%S{timezone_offset}")
        
        # signatureRandom: 16位随机字符串
        signature_random = ''.join(random.choices(string.ascii_letters + string.digits, k=16))
        
        # 获取音频时长（毫秒）
        duration_ms = _get_audio_duration(audio_path)
        
        # 构建请求参数（不包括signature）
        params = {
            'appId': appid,
            'accessKeyId': api_key,  # APIKey 作为 accessKeyId
            'dateTime': date_time,
            'signatureRandom': signature_random,
            'fileSize': str(file_len),
            'fileName': file_name,
            'duration': str(duration_ms),  # 音频时长（毫秒）
            'language': 'autodialect'  # 自动方言识别
        }
        
        # 生成签名：排除signature字段，排序，URL编码，HMAC-SHA1，Base64
        sorted_params = {}
        for k, v in sorted(params.items()):
            if k != 'signature' and v:
                sorted_params[k] = v
        
        # 构建 baseString
        parts = []
        for k, v in sorted_params.items():
            encoded_key = urllib.parse.quote(k, safe='')
            encoded_value = urllib.parse.quote(str(v), safe='')
            parts.append(f"{encoded_key}={encoded_value}")
        base_string = '&'.join(parts)
        
        # HMAC-SHA1
        signature_bytes = hmac.new(
            api_secret.encode('utf-8'),
            base_string.encode('utf-8'),
            hashlib.sha1
        ).digest()
        
        # Base64编码
        signature = base64.b64encode(signature_bytes).decode('utf-8')
        
        # 根据官方文档，signature 应该在请求头中，而不是 URL 参数
        # 构建上传 URL（参数在 URL 中，但不包括 signature）
        upload_params = {k: v for k, v in params.items() if k != 'signature'}
        upload_url = api_host + api_upload + "?" + urllib.parse.urlencode(upload_params, quote_via=urllib.parse.quote)
        
        # 第一步：上传文件
        _set_job(job_id, message="上传音频到科大讯飞（新版API）…", progress=20)
        _set_job(job_id, log_tail=[f"上传 URL: {upload_url[:200]}..."])
        _set_job(job_id, log_tail=[f"签名前10位: {signature[:10]}..."])
        
        with open(audio_path, 'rb') as f:
            audio_data = f.read()
        
        # 根据官方文档，signature 在请求头中
        upload_resp = requests.post(
            url=upload_url,
            headers={
                "Content-Type": "application/octet-stream",
                "signature": signature  # 签名在请求头中
            },
            data=audio_data,
            timeout=60
        )
        
        _set_job(job_id, log_tail=[f"上传响应状态码: {upload_resp.status_code}"])
        _set_job(job_id, log_tail=[f"上传响应: {upload_resp.text[:500]}"])
        
        if upload_resp.status_code != 200:
            return False, f"上传失败，状态码: {upload_resp.status_code}, 响应: {upload_resp.text[:500]}"
        
        try:
            upload_result = upload_resp.json()
        except Exception as e:
            return False, f"上传响应解析失败: {e}, 原始响应: {upload_resp.text[:500]}"
        
        # 检查错误
        # code: "000000" 表示成功，其他值表示失败
        code = upload_result.get('code')
        desc_info = upload_result.get('descInfo', '')
        
        # 成功的情况：code 为 "000000" 或 "0"，或者 descInfo 包含 "success"
        is_success = (
            code == "000000" or 
            code == "0" or 
            code == 0 or 
            "success" in desc_info.lower()
        )
        
        if not is_success:
            desc = upload_result.get('message') or upload_result.get('descInfo') or '未知错误'
            return False, f"上传失败: {desc}, code: {code}, 完整响应: {json.dumps(upload_result, ensure_ascii=False)[:500]}"
        
        # 获取订单ID
        order_id = upload_result.get('content', {}).get('orderId') or upload_result.get('data', {}).get('orderId') or upload_result.get('orderId')
        if not order_id:
            return False, f"上传响应中未找到 orderId: {upload_result}"
        
        # 第二步：轮询查询结果
        _set_job(job_id, message="等待转写结果…", progress=50)
        
        # 根据 taskEstimateTime 估算最大等待时间（毫秒转秒，再加一些缓冲）
        task_estimate_time = upload_result.get('content', {}).get('taskEstimateTime', 60000)  # 默认60秒
        max_wait_seconds = max(600, int(task_estimate_time / 1000) * 3)  # 至少10分钟，或预估时间的3倍
        max_polls = max(120, max_wait_seconds // 5)  # 每5秒查询一次，最多轮询 max_polls 次
        _set_job(job_id, log_tail=[f"预估处理时间: {task_estimate_time}ms ({task_estimate_time/1000:.1f}秒), 最大等待: {max_wait_seconds}秒, 最多查询: {max_polls}次"])
        
        poll_count = 0
        
        while poll_count < max_polls:
            poll_count += 1
            _set_job(job_id, message=f"查询转写结果… ({poll_count}/{max_polls})", progress=50 + min(40, poll_count * 40 // max_polls))
            
            # 每次查询都需要新的参数和签名
            now = datetime.now()
            timezone_offset = now.strftime('%z') or '+0800'
            # 确保格式为 ±HHmm（例如 +0800）
            if len(timezone_offset) == 5:
                # 已经是 ±HHmm 格式
                pass
            elif len(timezone_offset) == 6:
                # 可能是 ±HH:MM 格式，需要去掉冒号
                timezone_offset = timezone_offset[:3] + timezone_offset[4:]
            date_time = now.strftime(f"%Y-%m-%dT%H:%M:%S{timezone_offset}")
            signature_random = ''.join(random.choices(string.ascii_letters + string.digits, k=16))
            
            query_params = {
                'appId': appid,
                'accessKeyId': api_key,
                'dateTime': date_time,
                'signatureRandom': signature_random,
                'orderId': order_id
            }
            
            # 生成签名
            sorted_query_params = {}
            for k, v in sorted(query_params.items()):
                if k != 'signature' and v:
                    sorted_query_params[k] = v
            
            parts = []
            for k, v in sorted_query_params.items():
                encoded_key = urllib.parse.quote(k, safe='')
                encoded_value = urllib.parse.quote(str(v), safe='')
                parts.append(f"{encoded_key}={encoded_value}")
            base_string = '&'.join(parts)
            
            signature_bytes = hmac.new(
                api_secret.encode('utf-8'),
                base_string.encode('utf-8'),
                hashlib.sha1
            ).digest()
            signature = base64.b64encode(signature_bytes).decode('utf-8')
            
            # 根据官方文档，signature 应该在请求头中
            result_url = api_host + api_get_result + "?" + urllib.parse.urlencode(query_params, quote_via=urllib.parse.quote)
            result_resp = requests.get(
                url=result_url,
                headers={"signature": signature},  # 签名在请求头中
                timeout=30
            )
            
            if result_resp.status_code != 200:
                return False, f"查询失败，状态码: {result_resp.status_code}, 响应: {result_resp.text[:500]}"
            
            try:
                result_data = result_resp.json()
            except Exception as e:
                return False, f"查询响应解析失败: {e}, 原始响应: {result_resp.text[:500]}"
            
            # 检查错误
            # code: "000000" 表示成功，其他值表示失败
            code = result_data.get('code')
            desc_info = result_data.get('descInfo', '')
            
            # 成功的情况：code 为 "000000" 或 "0"，或者 descInfo 包含 "success"
            is_success = (
                code == "000000" or 
                code == "0" or 
                code == 0 or 
                "success" in desc_info.lower()
            )
            
            if not is_success:
                desc = result_data.get('message') or result_data.get('descInfo') or '未知错误'
                return False, f"查询失败: {desc}, code: {code}, 完整响应: {json.dumps(result_data, ensure_ascii=False)[:500]}"
            
            # 检查状态
            # 根据响应，status 在 content.orderInfo.status 中
            # status: 0或3=处理中, 4=完成, 2=失败
            order_info = result_data.get('content', {}).get('orderInfo', {})
            status = order_info.get('status') or result_data.get('content', {}).get('status') or result_data.get('data', {}).get('status') or result_data.get('status')
            
            _set_job(job_id, log_tail=[f"查询状态: {status}, orderInfo: {json.dumps(order_info, ensure_ascii=False)[:200]}"])
            
            # status: 0或3=处理中, 4=完成, 2=失败
            if status == 4:
                # 完成，提取文本结果
                result_text = ""
                # 尝试从 content.orderResult 或 content 中获取结果
                content = result_data.get('content', {})
                order_result = content.get('orderResult', '')
                
                # 优先使用 orderResult
                if order_result:
                    if isinstance(order_result, str):
                        # 如果是字符串，尝试解析 JSON
                        try:
                            order_result = json.loads(order_result)
                        except:
                            pass
                    
                    if isinstance(order_result, dict):
                        # 尝试从 lattice -> json_1best 结构中提取文本
                        lattice = order_result.get('lattice', [])
                        if lattice and len(lattice) > 0:
                            # 遍历所有 lattice 项（可能有多段）
                            all_words = []
                            for lattice_item in lattice:
                                json_1best_str = lattice_item.get('json_1best', '')
                                if json_1best_str:
                                    try:
                                        json_1best = json.loads(json_1best_str)
                                        # 从 st -> rt -> ws -> cw -> w 中提取文本
                                        st = json_1best.get('st', {})
                                        rt = st.get('rt', [])
                                        
                                        # 遍历所有 rt 项（每段可能有多个 rt）
                                        for rt_item in rt:
                                            ws = rt_item.get('ws', [])
                                            for ws_item in ws:
                                                cw = ws_item.get('cw', [])
                                                for cw_item in cw:
                                                    word = cw_item.get('w', '')
                                                    if word:
                                                        all_words.append(word)
                                    except Exception as e:
                                        _set_job(job_id, log_tail=[f"解析 json_1best 失败: {e}, 内容: {json_1best_str[:200]}"])
                            
                            if all_words:
                                result_text = ''.join(all_words)
                                _set_job(job_id, log_tail=[f"提取到 {len(all_words)} 个词，总长度: {len(result_text)} 字符"])
                        
                        # 如果 lattice 解析失败，尝试其他字段
                        if not result_text:
                            result_text = order_result.get('text', order_result.get('content', ''))
                    
                    elif isinstance(order_result, list):
                        texts = []
                        for item in order_result:
                            if isinstance(item, dict):
                                texts.append(item.get('text', item.get('content', str(item))))
                            else:
                                texts.append(str(item))
                        result_text = '\n'.join(filter(None, texts))
                
                # 如果 orderResult 为空，尝试其他字段
                if not result_text:
                    data = result_data.get('data', {}) or content
                    
                    if 'result' in data:
                        result_field = data['result']
                        if isinstance(result_field, str):
                            result_text = result_field
                        elif isinstance(result_field, dict):
                            result_text = result_field.get('text', result_field.get('content', ''))
                        elif isinstance(result_field, list):
                            texts = []
                            for item in result_field:
                                if isinstance(item, dict):
                                    texts.append(item.get('text', item.get('content', str(item))))
                                else:
                                    texts.append(str(item))
                            result_text = '\n'.join(filter(None, texts))
                    
                    if 'text' in data:
                        result_text = data['text']
                
                if result_text:
                    return True, result_text.strip()
                else:
                    return False, f"转写完成但未找到文本结果。响应: {json.dumps(result_data, ensure_ascii=False)[:500]}"
            
            if status == 2:
                return False, f"转写失败: {result_data.get('message', '未知错误')}"
            
            # status == 0 或 3，继续等待（处理中）
            # 根据 taskEstimateTime 调整等待时间
            wait_time = 5  # 默认5秒
            task_estimate = result_data.get('content', {}).get('taskEstimateTime', 0)
            if task_estimate > 0:
                # 如果预估时间还很长，可以稍微延长等待间隔
                remaining_estimate = task_estimate / 1000  # 转为秒
                if remaining_estimate > 60:
                    wait_time = 10  # 如果预估时间超过1分钟，每10秒查询一次
                elif remaining_estimate > 30:
                    wait_time = 7   # 如果预估时间超过30秒，每7秒查询一次
            
            time.sleep(wait_time)  # 等待后再次查询
        
        return False, f"转写超时（已查询 {poll_count} 次）"
        
    except requests.exceptions.RequestException as e:
        return False, f"网络请求异常: {e}"
    except Exception as e:
        return False, f"新版 API 调用异常: {e}"


def _xunfei_transcribe(audio_path: Path, job_id: str, appid: str, secret_key: str) -> tuple[bool, str]:
    """
    科大讯飞 HTTP API 转写（raasr.xfyun.cn）
    根据用户提供的 demo 代码实现
    返回 (成功, 文本或错误信息)
    """
    if requests is None:
        return False, "需要安装 requests 库：pip install requests"

    lfasr_host = 'https://raasr.xfyun.cn/v2/api'
    api_upload = '/upload'
    api_get_result = '/getResult'

    try:
        ts = str(int(time.time()))
        signa = _xunfei_get_signa(appid, secret_key, ts)

        # 第一步：上传文件
        _set_job(job_id, message="上传音频到科大讯飞…", progress=20)
        file_len = audio_path.stat().st_size
        file_name = audio_path.name

        param_dict = {
            'appId': appid,
            'signa': signa,
            'ts': ts,
            'fileSize': file_len,
            'fileName': file_name,
            'duration': '200'  # 默认值，实际可能需要根据音频时长计算
        }
        
        # 调试：记录签名信息（不包含完整密钥）
        _set_job(job_id, log_tail=[f"签名调试: appId={appid}, ts={ts}, signa前10位={signa[:10]}..."])

        with open(audio_path, 'rb') as f:
            audio_data = f.read()

        # 构建上传 URL
        upload_url = lfasr_host + api_upload + "?" + urllib.parse.urlencode(param_dict, quote_via=urllib.parse.quote)
        
        # 调试：记录完整 URL（不包含音频数据）
        _set_job(job_id, log_tail=[f"上传 URL: {upload_url[:200]}..."])
        
        upload_resp = requests.post(
            url=upload_url,
            headers={"Content-type": "application/json"},
            data=audio_data,
            timeout=60
        )

        # 详细日志：记录请求和响应
        log_lines = [
            f"上传 URL: {upload_url[:200]}...",
            f"HTTP 状态码: {upload_resp.status_code}",
            f"响应头: {dict(upload_resp.headers)}",
            f"响应文本前500字符: {upload_resp.text[:500]}"
        ]
        _set_job(job_id, log_tail=log_lines)

        if upload_resp.status_code != 200:
            return False, f"上传失败，状态码: {upload_resp.status_code}, 响应: {upload_resp.text[:500]}"

        try:
            upload_result = upload_resp.json()
        except Exception as e:
            return False, f"上传响应解析失败: {e}, 原始响应: {upload_resp.text[:500]}"

        # 详细日志：记录解析后的响应
        log_lines.append(f"解析后的响应: {json.dumps(upload_result, ensure_ascii=False)[:500]}")
        _set_job(job_id, log_tail=log_lines)

        # code 可能是字符串或数字，检查是否为成功（0 或 "0"）
        code = upload_result.get('code')
        if code != 0 and str(code) != "0":
            desc = upload_result.get('desc') or upload_result.get('descInfo') or '未知错误'
            log_lines.append(f"错误: {desc}, code: {code}")
            _set_job(job_id, log_tail=log_lines)
            return False, f"上传失败: {desc}, code: {code}, 完整响应: {json.dumps(upload_result, ensure_ascii=False)[:500]}"

        order_id = upload_result.get('content', {}).get('orderId')
        if not order_id:
            return False, f"上传响应中未找到 orderId: {upload_result}"

        # 第二步：轮询查询结果
        _set_job(job_id, message="等待转写结果…", progress=50)
        
        max_polls = 120  # 最多轮询 120 次（10分钟）
        poll_count = 0
        status = 3  # 3=处理中，4=完成

        while status == 3 and poll_count < max_polls:
            poll_count += 1
            _set_job(job_id, message=f"查询转写结果… ({poll_count}/{max_polls})", progress=50 + min(40, poll_count * 40 // max_polls))

            # 每次查询都需要新的 ts 和 signa
            ts_query = str(int(time.time()))
            signa_query = _xunfei_get_signa(appid, secret_key, ts_query)
            param_dict = {
                'appId': appid,
                'signa': signa_query,
                'ts': ts_query,
                'orderId': order_id,
                'resultType': 'transfer,predict'
            }

            result_url = lfasr_host + api_get_result + "?" + urllib.parse.urlencode(param_dict)
            result_resp = requests.post(
                url=result_url,
                headers={"Content-type": "application/json"},
                timeout=30
            )

            if result_resp.status_code != 200:
                return False, f"查询失败，状态码: {result_resp.status_code}, 响应: {result_resp.text[:500]}"

            try:
                result_data = result_resp.json()
            except Exception as e:
                return False, f"查询响应解析失败: {e}, 原始响应: {result_resp.text[:500]}"

            # code 可能是字符串或数字，检查是否为成功（0 或 "0"）
            code = result_data.get('code')
            if code != 0 and str(code) != "0":
                desc = result_data.get('desc') or result_data.get('descInfo') or '未知错误'
                return False, f"查询失败: {desc}, code: {code}, 完整响应: {json.dumps(result_data, ensure_ascii=False)[:500]}"

            order_info = result_data.get('content', {}).get('orderInfo', {})
            status = order_info.get('status', 3)

            if status == 4:
                # 完成，提取文本结果
                result_text = ""
                content = result_data.get('content', {})
                order_info = content.get('orderInfo', {})
                
                # 尝试多种可能的字段路径
                # 1. content.orderInfo.result (常见格式)
                if 'result' in order_info:
                    result_data_field = order_info['result']
                    if isinstance(result_data_field, str):
                        result_text = result_data_field
                    elif isinstance(result_data_field, list):
                        # 分段结果，合并
                        texts = []
                        for item in result_data_field:
                            if isinstance(item, dict):
                                # 可能的结构：{"oneBest": "文本", "speaker": 1} 或 {"text": "文本"}
                                text = item.get('oneBest') or item.get('text') or item.get('content') or str(item)
                                if text:
                                    texts.append(text)
                            else:
                                texts.append(str(item))
                        result_text = '\n'.join(filter(None, texts))
                    elif isinstance(result_data_field, dict):
                        # 可能是嵌套结构
                        result_text = json.dumps(result_data_field, ensure_ascii=False)
                
                # 2. content.result
                if not result_text and 'result' in content:
                    result_data_field = content['result']
                    if isinstance(result_data_field, str):
                        result_text = result_data_field
                    elif isinstance(result_data_field, list):
                        texts = []
                        for item in result_data_field:
                            if isinstance(item, dict):
                                text = item.get('oneBest') or item.get('text') or item.get('content') or str(item)
                                if text:
                                    texts.append(text)
                            else:
                                texts.append(str(item))
                        result_text = '\n'.join(filter(None, texts))
                
                # 3. content.text
                if not result_text and 'text' in content:
                    result_text = str(content['text'])
                
                # 4. content.data
                if not result_text and 'data' in content:
                    data = content['data']
                    if isinstance(data, list) and len(data) > 0:
                        texts = []
                        for item in data:
                            if isinstance(item, dict):
                                text = item.get('oneBest') or item.get('text') or item.get('content') or str(item)
                                if text:
                                    texts.append(text)
                            else:
                                texts.append(str(item))
                        result_text = '\n'.join(filter(None, texts))
                    elif isinstance(data, str):
                        result_text = data
                
                # 5. 如果还是找不到，返回完整响应用于调试
                if not result_text.strip():
                    full_response = json.dumps(result_data, ensure_ascii=False, indent=2)
                    _set_job(job_id, log_tail=[f"完整响应: {full_response[:1000]}"])
                    return False, f"转写完成但未找到文本结果。响应结构已记录到日志，请查看任务详情或联系开发者。响应预览: {full_response[:500]}"

                return True, result_text.strip()

            time.sleep(5)  # 等待 5 秒后再次查询

        return False, f"转写超时（已查询 {poll_count} 次），最后状态: {status}"

    except requests.exceptions.RequestException as e:
        return False, f"网络请求异常: {e}"
    except Exception as e:
        return False, f"科大讯飞 API 调用异常: {e}"


def _api_transcribe(audio_path: Path, job_id: str) -> tuple[bool, str]:
    """
    调用外部 STT API 进行转写。
    返回 (成功, 文本或错误信息)
    """
    cfg = _load_local_config()
    api_type = cfg.get("stt_api_type", "").strip().lower()

    # 科大讯飞 API
    if api_type == "xunfei":
        appid = cfg.get("stt_api_appid", "").strip()
        api_key = cfg.get("stt_api_key", "").strip()
        api_secret = cfg.get("stt_api_secret_key", "").strip()
        
        if not appid:
            return False, "未配置科大讯飞 appid（请在 config.json 中设置 stt_api_appid）"
        
        if not api_key and not api_secret:
            return False, "未配置科大讯飞 APIKey 或 APISecret（请在 config.json 中设置 stt_api_key 或 stt_api_secret_key）"
        
        # 如果同时配置了 APIKey 和 APISecret，优先尝试新版 HTTPS API（官方文档中的接口）
        if api_key and api_secret:
            _set_job(job_id, log_tail=[f"尝试使用新版 HTTPS API (office-api-ist-dx.iflyaisol.com)"])
            result = _xunfei_transcribe_new_api(audio_path, job_id, appid, api_key, api_secret)
            # 如果新版 API 成功或返回明确的错误（不是"暂未实现"），则返回结果
            if result[0] or ("暂未实现" not in result[1] and "回退" not in result[1]):
                return result
        
        # 回退到 HTTP API (raasr.xfyun.cn)
        _set_job(job_id, log_tail=[f"使用 HTTP API (raasr.xfyun.cn)"])
        # 优先使用 APIKey，如果没有则使用 APISecret
        secret_key = api_key if api_key else api_secret
        if api_key:
            _set_job(job_id, log_tail=[f"使用 APIKey 生成签名"])
        else:
            _set_job(job_id, log_tail=[f"使用 APISecret 生成签名"])
        
        return _xunfei_transcribe(audio_path, job_id, appid, secret_key)

    # 通用 API（原有逻辑）
    api_url = cfg.get("stt_api_url", "").strip()
    api_key = cfg.get("stt_api_key", "").strip()
    api_method = cfg.get("stt_api_method", "POST").strip().upper()

    if not api_url:
        return False, "未配置 STT API（请在 config.json 中设置 stt_api_url 或 stt_api_type=xunfei）"

    try:
        _set_job(job_id, message="调用 STT API 转写中…", progress=50)

        # 读取音频文件
        audio_data = audio_path.read_bytes()

        # 构建请求
        req = urllib.request.Request(api_url, data=audio_data)
        req.add_header("Content-Type", "audio/wav")
        if api_key:
            req.add_header("Authorization", f"Bearer {api_key}")

        # 发送请求
        with urllib.request.urlopen(req, timeout=300) as resp:
            if resp.status != 200:
                return False, f"STT API 返回错误状态码: {resp.status}"
            result_text = resp.read().decode("utf-8", errors="ignore").strip()

        if not result_text:
            return False, "STT API 返回空结果"

        return True, result_text

    except urllib.error.HTTPError as e:
        error_body = e.read().decode("utf-8", errors="ignore")[:500]
        return False, f"STT API HTTP 错误 {e.code}: {error_body}"
    except urllib.error.URLError as e:
        return False, f"STT API 网络错误: {e.reason}"
    except Exception as e:
        return False, f"STT API 调用异常: {e}"


def _worker(job_id: str, src_path: Path, mode: str = "local"):
    """
    转写工作线程。
    mode: "local" 使用本地 whisper-cli，"api" 使用外部 STT API
    """
    transcribe_start_time = time.time()
    _set_job(job_id, status="running", message="开始处理音频…", started_at=transcribe_start_time)
    try:
        if mode == "api":
            # API 模式：直接使用原始文件或转换为 WAV
            wav_path = WORK_DIR / f"{job_id}.wav"
            _set_job(job_id, message="转码中（ffmpeg）…", progress=10)
            ok, ffmpeg_log = _to_wav_16k_mono(src_path, wav_path)
            if not ok:
                _set_job(job_id, status="error", message="ffmpeg 转换失败（请确认已安装 ffmpeg）", log=ffmpeg_log)
                return

            _set_job(job_id, message="调用 STT API 转写中…", progress=30)
            ok, result = _api_transcribe(wav_path, job_id)
            if not ok:
                _set_job(job_id, status="error", message=f"STT API 转写失败：{result}")
                return

            text = result
            transcribe_end_time = time.time()
            transcribe_duration = transcribe_end_time - transcribe_start_time
            _set_job(job_id, status="done", message="完成", text=text, finished_at=transcribe_end_time, transcribe_duration=transcribe_duration)

        else:
            # 本地模式（默认）：使用 whisper-cli
            if not Path(WHISPER_MODEL).exists():
                _set_job(
                    job_id,
                    status="error",
                    message=f"模型文件不存在：{WHISPER_MODEL}（请下载 ggml 模型并放到 models/ 目录）",
                )
                return

            wav_path = WORK_DIR / f"{job_id}.wav"
            _set_job(job_id, message="转码中（ffmpeg）…", progress=0)
            ok, ffmpeg_log = _to_wav_16k_mono(src_path, wav_path)
            if not ok:
                _set_job(job_id, status="error", message="ffmpeg 转换失败（请确认已安装 ffmpeg）", log=ffmpeg_log)
                return

            _set_job(job_id, status="running", message="开始转写（Whisper）…", progress=0)
            out_prefix = WORK_DIR / f"{job_id}"
            ok, whisper_log = _whisper_transcribe(wav_path, out_prefix)
            if not ok:
                _set_job(job_id, status="error", message="Whisper 转写失败（请确认 whisper 可执行文件可用）", log=whisper_log)
                return

            txt_path = WORK_DIR / f"{job_id}.txt"
            if not txt_path.exists():
                # 有的版本会生成 out_prefix + ".txt"
                alt = Path(str(out_prefix) + ".txt")
                if alt.exists():
                    txt_path = alt

            text = txt_path.read_text(encoding="utf-8", errors="ignore") if txt_path.exists() else ""
            transcribe_end_time = time.time()
            transcribe_duration = transcribe_end_time - transcribe_start_time
            _set_job(job_id, status="done", message="完成", text=text, finished_at=transcribe_end_time, transcribe_duration=transcribe_duration, log=whisper_log)

        # 额外：在 survey/ 目录落一份结果，方便你在"访谈材料目录"直接看到输出
        original_name = (_get_job(job_id) or {}).get("original_filename") or f"{job_id}{src_path.suffix}"
        out_name = f"{_safe_basename(original_name)}_{job_id}.txt"
        (SURVEY_DIR / out_name).write_text(text, encoding="utf-8")
    except Exception as e:
        _set_job(job_id, status="error", message=f"服务异常：{e}")


@app.get("/")
def index():
    return send_from_directory(str(ROOT_DIR), "index.html")


@app.get("/README.md")
def readme():
    return send_from_directory(str(ROOT_DIR), "README.md")


@app.post("/api/transcribe")
def transcribe():
    # 获取转写模式：local（本地 whisper-cli）或 api（外部 STT API）
    mode = request.form.get("mode", "local").strip().lower()
    if mode not in ["local", "api"]:
        return jsonify({"error": f"无效的 mode 参数：{mode}，必须是 'local' 或 'api'"}), 400

    if mode == "local":
        # 本地模式需要检查模型文件
        if not Path(WHISPER_MODEL).exists():
            return (
                jsonify(
                    {
                        "error": f"模型文件不存在：{WHISPER_MODEL}。请下载 ggml 模型并放到项目根目录 models/ 下（例如 models/ggml-small.bin）。"
                    }
                ),
                400,
            )
    else:
        # API 模式需要检查配置
        cfg = _load_local_config()
        api_type = cfg.get("stt_api_type", "").strip().lower()
        
        if api_type == "xunfei":
            # 科大讯飞 API：需要 appid 和 secret_key（或 APIKey）
            appid = cfg.get("stt_api_appid", "").strip()
            api_key = cfg.get("stt_api_key", "").strip()
            secret_key = cfg.get("stt_api_secret_key", "").strip()
            if not appid or (not api_key and not secret_key):
                return jsonify({"error": "未配置科大讯飞 API（请在 config.json 中设置 stt_api_appid 和 stt_api_key 或 stt_api_secret_key）"}), 400
        else:
            # 通用 API：需要 URL
            if not cfg.get("stt_api_url", "").strip():
                return jsonify({"error": "未配置 STT API（请在 config.json 中设置 stt_api_url 或 stt_api_type=xunfei）"}), 400

    if "file" not in request.files:
        return jsonify({"error": "缺少文件字段 file"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "文件名为空"}), 400

    job_id = uuid.uuid4().hex
    suffix = Path(f.filename).suffix or ".audio"
    src_path = UPLOAD_DIR / f"{job_id}{suffix}"
    f.save(str(src_path))

    _set_job(
        job_id,
        status="queued",
        message="已接收，排队中…",
        created_at=time.time(),
        original_filename=f.filename,
    )
    t = threading.Thread(target=_worker, args=(job_id, src_path, mode), daemon=True)
    t.start()

    return jsonify({"job_id": job_id})


@app.get("/api/jobs/<job_id>")
def job(job_id: str):
    j = _get_job(job_id)
    if not j:
        return jsonify({"error": "job 不存在"}), 404
    # 只返回必要字段，避免日志过大
    resp = {
        "job_id": job_id,
        "status": j.get("status"),
        "message": j.get("message"),
        "progress": j.get("progress"),
        "text": j.get("text", ""),
        "log_tail": j.get("log_tail", []),
        "transcribe_duration": j.get("transcribe_duration"),
    }
    return jsonify(resp)


@app.get("/api/jobs")
def list_jobs():
    """列出最近的任务（最多20个），用于调试"""
    with _jobs_lock:
        jobs_list = list(_jobs.items())
        # 按创建时间倒序排列
        jobs_list.sort(key=lambda x: x[1].get("created_at", 0) or 0, reverse=True)
        result = []
        for job_id, job_data in jobs_list[:20]:
            result.append({
                "job_id": job_id,
                "status": job_data.get("status"),
                "message": job_data.get("message"),
                "created_at": job_data.get("created_at"),
                "finished_at": job_data.get("finished_at"),
            })
        return jsonify({"jobs": result})


@app.get("/api/jobs/<job_id>/download")
def download_job_text(job_id: str):
    """
    下载转写文本（.txt）。优先使用内存里的 text；否则读取 data/work/<job_id>.txt
    """
    j = _get_job(job_id)
    if j and j.get("text"):
        tmp = WORK_DIR / f"{job_id}.txt"
        tmp.write_text(j.get("text", ""), encoding="utf-8")
        return send_file(
            str(tmp),
            mimetype="text/plain; charset=utf-8",
            as_attachment=True,
            download_name=f"{job_id}.txt",
        )

    txt_path = WORK_DIR / f"{job_id}.txt"
    if not txt_path.exists():
        alt = WORK_DIR / f"{job_id}.txt"
        if not alt.exists():
            return jsonify({"error": "转写文本不存在或任务未完成"}), 404
        txt_path = alt

    return send_file(
        str(txt_path),
        mimetype="text/plain; charset=utf-8",
        as_attachment=True,
        download_name=f"{job_id}.txt",
    )


def _extract_questions_from_docx(docx_path: Path) -> list[dict]:
    """
    从 docx 里提取“编号清晰”的问题。
    输出: [{id: '1', text: '...'}, ...]
    """
    doc = Document(str(docx_path))
    questions: list[dict] = []

    # 常见编号：1. / 1、/ 1) / （1）/ 1）/ Q1 / Q1:
    q_re = re.compile(r"^(?:Q\\s*)?(\\d{1,3})\\s*(?:[\\.、:：\\)）]|（\\1）)\\s*(.+)$", re.IGNORECASE)

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        m = q_re.match(t)
        if not m:
            continue
        qid = m.group(1).strip()
        qtext = m.group(2).strip()
        if qtext:
            questions.append({"id": qid, "text": qtext})

    return questions


def _extract_full_text_from_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines)


def _build_qa_prompt(transcript: str, questions_text: str) -> str:
    # 目标：输出“问题/答案”的干净纯文本（用户已验证的提示词风格）
    return f"""我上传了两份文件，一份是录音.txt，是对采访者的录音内容。 questions.txt 这是准备好的问题，我需要你分析录音的内容，并把里面的内容分别匹配到对应的 questions 的问题里面，但是录音中，无法区分出采访者和被采访者，你只能自己去识别判断。

请严格按 questions.txt 的分类标题与题目顺序输出。输出只允许包含“分类标题 + 问题 + 答案”，不要输出其它任何干扰内容（不要解释、不要规则、不要 JSON、不要 Markdown、不要代码块）。
并且**每个问题必须带题号**，格式为：`题号. 问题原文`（例如：`1. 最初发现...？`）。
重要：你必须覆盖 questions.txt 中的**四大类全部问题**，不得只输出第一类/前半部分。

输出格式模板（必须严格遵守）：
三、学龄前康复阶段（17 题）

1. 最初发现孩子可能存在发育异常的人是谁？
录音内容：老人带得多，语言出得慢，关注物体（轮子、风扇），带出去时被亲友提醒。

2. 从孩子最初被怀疑异常，到最终被确诊孤独症，整个过程用了多长时间？
录音中未提及确诊时长，仅提及发现过程。

（后续每题同样格式：问题一行 + 下一行以“录音内容：...”或“录音中未提及...”开头）

【questions.txt】：
{questions_text}

【录音.txt】：
{transcript}

现在开始输出最终结果（只输出结果正文）："""


def _openrouter_chat(api_key: str, model: str, prompt: str, max_tokens: int) -> dict:
    """
    OpenRouter（OpenAI 兼容）接口：
    POST https://openrouter.ai/api/v1/chat/completions
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是一个擅长定性访谈分析与信息抽取的助手。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url=url,
        data=data,
        method="POST",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            # OpenRouter 推荐带上这两个 header（可选）
            "HTTP-Referer": "http://127.0.0.1:8000",
            "X-Title": "Local Survey Tool",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
            return json.loads(raw)
    except urllib.error.HTTPError as e:
        raw = e.read().decode("utf-8", errors="ignore") if hasattr(e, "read") else str(e)
        raise RuntimeError(f"OpenRouter HTTPError: {e.code} {raw}")
    except urllib.error.URLError as e:
        raise RuntimeError(f"OpenRouter URLError: {e}")


@app.post("/api/bundle")
def make_bundle():
    """
    将“问题清单(docx)” + “转写文本(job_id)” 打包成 JSON，方便用户调用自己的 LLM API。
    不在服务端调用任何大模型。
    """
    job_id = (request.form.get("job_id") or "").strip()
    if not job_id:
        return jsonify({"error": "缺少 job_id"}), 400

    txt_path = WORK_DIR / f"{job_id}.txt"
    if not txt_path.exists():
        return jsonify({"error": f"找不到转写文本：{txt_path}（请确认任务已完成）"}), 404

    if "docx" not in request.files:
        return jsonify({"error": "缺少 docx 文件字段 docx"}), 400

    f = request.files["docx"]
    if not f.filename or not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "请上传 .docx 文件"}), 400

    bundle_id = uuid.uuid4().hex
    docx_path = WORK_DIR / f"{bundle_id}.docx"
    f.save(str(docx_path))

    try:
        questions = _extract_questions_from_docx(docx_path)
        docx_text = _extract_full_text_from_docx(docx_path)
    except Exception as e:
        return jsonify({"error": f"解析 docx 失败：{e}"}), 400

    transcript = txt_path.read_text(encoding="utf-8", errors="ignore")
    payload = {
        "job_id": job_id,
        "docx_name": f.filename,
        # 你可以把 docx_text 直接喂给你的大模型（比我用正则抽题更“语义完整”）
        "docx_text": docx_text,
        # 同时保留我抽取的结构化 questions，用于约束/校验（可选使用）
        "questions": questions,
        "transcript": transcript,
    }

    out_path = WORK_DIR / f"bundle-{job_id}.json"
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return send_file(
        str(out_path),
        mimetype="application/json; charset=utf-8",
        as_attachment=True,
        download_name=f"bundle-{job_id}.json",
    )


@app.post("/api/llm/match")
def llm_match():
    """
    通过 OpenRouter 调用大模型做“问题-答案匹配”。
    注意：不保存 key，不做持久化，只做一次请求。
    """
    body = request.get_json(silent=True) or {}
    transcript = (body.get("transcript") or "").strip()
    questions = (body.get("questions") or "").strip()

    cfg = _load_local_config()
    api_key = cfg.get("openrouter_api_key", "")
    model = cfg.get("openrouter_model", "tngtech/deepseek-r1t2-chimera:free")
    max_tokens = int(cfg.get("openrouter_max_tokens", 8192))

    if not api_key:
        return jsonify({"error": "未配置 OpenRouter API Key：请在项目根目录创建 config.json（参考 config.example.json）"}), 400
    if not transcript:
        return jsonify({"error": "缺少 transcript（录音转写文本）"}), 400
    if not questions:
        return jsonify({"error": "缺少 questions（问题模板文本）"}), 400

    prompt = _build_qa_prompt(transcript=transcript, questions_text=questions)
    match_start_time = time.time()
    try:
        resp = _openrouter_chat(api_key=api_key, model=model, prompt=prompt, max_tokens=max_tokens)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    match_end_time = time.time()
    match_duration = match_end_time - match_start_time

    content = ""
    finish_reason = ""
    usage = {}
    try:
        content = resp["choices"][0]["message"]["content"]
        finish_reason = resp["choices"][0].get("finish_reason", "")
    except Exception:
        content = ""

    cleaned = _strip_code_fence(content)
    try:
        usage = resp.get("usage") or {}
    except Exception:
        usage = {}

    return jsonify(
        {
            "model": model,
            "max_tokens": max_tokens,
            "finish_reason": finish_reason,
            "usage": usage,
            "cleaned": cleaned,
            "match_duration": match_duration,
        }
    )

@app.get("/api/health")
def health():
    return jsonify(
        {
            "ok": True,
            "whisper_bin": WHISPER_BIN,
            "whisper_bin_path": shutil.which(WHISPER_BIN) if WHISPER_BIN else None,
            "ffmpeg_bin": FFMPEG_BIN,
            "ffmpeg_bin_path": shutil.which(FFMPEG_BIN) if FFMPEG_BIN else None,
            "model": WHISPER_MODEL,
            "model_exists": Path(WHISPER_MODEL).exists(),
            "language": WHISPER_LANGUAGE,
            "max_mb": MAX_CONTENT_LENGTH_MB,
        }
    )


if __name__ == "__main__":
    # 默认只监听本机，云端部署时可通过环境变量打开外网监听
    host = _env("HOST", "127.0.0.1")
    port = int(_env("PORT", "8000"))
    debug = _env("DEBUG", "1").strip() in ("1", "true", "True", "yes", "YES")
    app.run(host=host, port=port, debug=debug, threaded=True)

